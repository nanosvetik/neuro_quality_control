import os
import re
from datetime import datetime
import json
import ffmpeg
import torch
import whisperx
import logging
import librosa
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import zipfile
import lmstudio as lms
from transformers import Wav2Vec2ForSequenceClassification, Wav2Vec2FeatureExtractor
from sentence_transformers import SentenceTransformer
import faiss

logging.basicConfig(level=logging.INFO)

OUTPUT_DIR = '.'
TEMP_DIR = './temp'
ENCODING_TEXT = 'utf-8'
FAISS_INDEX = './call-center-knowledge-base/vector_index.faiss'
FAISS_METADATA = './call-center-knowledge-base/vector_metadata.json'

class Worker():

    def __init__(self):
        pass

    @staticmethod
    def get_system_prompt(criteria_text):
        # Обновленный промт с учетом базы знаний
        system_prompt = f'''****Ты специалист по оценке качества работы менеджера колл-центра. Задача:
        Проанализируй текст транскрипции звонка менеджера колл-центра. Учти, что текст может содержать ошибки из-за работы модели ASR (автоматического распознавания речи). На основе анализа сгенерируй два отчета:
        1. Отчет по оценке работы менеджера (с баллами и комментариями [обоснуй оценку, при необходимости приведи пример из текста]).
        2. Отчет по рекомендациям менеджеру (с конкретными предложениями по улучшению).
        Избегай общих фраз, приводи конкретные примеры из диалога, если это необходимо. Будь предельно внимательной и объективной, от твоей оценки зависит карьера менеджера и репутация компании.

        ### Важные указания:
        1. **Используй только шкалу оценки из базы знаний**. Не придумывай свои баллы или критерии.
        2. **Ошибки ASR не должны влиять на оценку**. Если в тексте есть ошибки распознавания речи, игнорируй их и оценивай только действия менеджера. В канале менеджера может присутсвовать речь клиента. внимательно проверяй и игнорируй эти выражения.
        3. **Если событие отсутствует, вместо оценки укажи "Не выявлено"**.

        ### Классификация типа звонка:
        Определи тип звонка на основе контекста. Возможные типы:
        Товар (консультация, использование, проблема), Заказ (информация, оформление, отмена), Доставка (проблема, информация), Возврат/Обмен, Адрес АСЦ, Закупка комплектующих, Программа лояльности, Юридическое лицо, Жалоба, Рассылка/реклама, Нецелевой звонок.

        ### Критерии и шкала оценки:
        {criteria_text}

        ### Формат отчетов:
        1. **Отчет по оценке работы менеджера:**
        - Тематика звонка: [Тема].
        - Краткая саммаризация разговора: [Краткое содержание].
        - Оценка по каждому критерию:
            - [Название критерия]: [Баллы] 
            - [Комментарий]. Пример: [Точный пример из текста с таймкодом].
        - Вывод: [Итоговая оценка работы менеджера с четким обоснованием].

        2. **Отчет по рекомендациям менеджеру:**
        - Рекомендации по улучшению работы менеджера.
        - План действий (обучение, практика, обратная связь).
        '''
        return system_prompt
    
    @staticmethod
    def get_user_message(transcript):
        user_message = f"Проанализируй текст транскрипции звонка менеджера колл-центра {transcript}."
        return user_message

    def process_audio_file(self, audio_file):

        def parse_filename(filename):
            # Обновлённое регулярное выражение
            pattern = r"(\d{2}\.\d{2}\.\d{4})__(\d{2}-\d{2}-\d{2})__(_*)__(\d+)"
            match = re.match(pattern, filename)
            if not match:
                raise ValueError(f"Название файла '{filename}' не соответствует ожидаемому формату.")
            
            date_str, time_str, employee_name, phone_number = match.groups()
            date = datetime.strptime(date_str, "%d.%m.%Y").date()
            time = datetime.strptime(time_str, "%H-%M-%S").time()
            
            # Если ФИ отсутствует (только подчёркивания), используем значение по умолчанию
            if employee_name.strip("_") == "":
                employee_name = "Неизвестный_Сотрудник"
            else:
                # Убираем лишние подчёркивания, если они есть
                employee_name = employee_name.strip("_")
            
            return {
                "date": date,
                "time": time,
                "employee_name": employee_name,
                "phone_number": phone_number
            }


        # Функция для получения длительности аудио
        def get_audio_duration(input_file):
            try:
                probe = ffmpeg.probe(input_file)
                duration = float(probe["format"]["duration"])  # Длительность в секундах
                return duration
            except Exception as e:
                logging.error(f"Ошибка при получении длительности аудио: {e}")
                return None
            
        # Функция для разделения текста на фразы
        def split_into_phrases(segment):
            text = segment["text"]
            start = segment["start"]
            end = segment["end"]
            
            # Разделяем текст на фразы по знакам препинания
            phrases = re.split(r"(?<=[.!?]) +", text)
            
            # Рассчитываем временные метки для каждой фразы
            phrase_segments = []
            total_duration = end - start
            phrase_duration = total_duration / len(phrases)
            
            for i, phrase in enumerate(phrases):
                phrase_start = start + i * phrase_duration
                phrase_end = phrase_start + phrase_duration
                phrase_segments.append({
                    "start": phrase_start,
                    "end": phrase_end,
                    "text": phrase.strip()
                })
            
            return phrase_segments

        # Функция для форматирования вывода транскрипции
        def format_transcription(segments, speaker_label):
            formatted_text = []
            for segment in segments:
                # Разделяем сегмент на фразы
                phrases = split_into_phrases(segment)
                for phrase in phrases:
                    start = phrase["start"]
                    end = phrase["end"]
                    text = phrase["text"]
                    formatted_text.append(f"[{start:.2f}-{end:.2f}] {speaker_label}: {text}")
            return "\n".join(formatted_text)

        # Функция для разделения аудио на каналы и предобработки
        def preprocess_audio(input_file, output_dir):
            os.makedirs(output_dir, exist_ok=True)
            
            left_channel = os.path.join(output_dir, "left_channel.wav")
            right_channel = os.path.join(output_dir, "right_channel.wav")
            
            try:
                # Разделяем аудио на два канала
                (
                    ffmpeg.input(input_file)
                    .output(left_channel, ac=1, af="pan=stereo|c0=c0", ar=16000)
                    .run(capture_stdout=True, capture_stderr=True, overwrite_output=True)
                )
                (
                    ffmpeg.input(input_file)
                    .output(right_channel, ac=1, af="pan=stereo|c1=c1", ar=16000)
                    .run(capture_stdout=True, capture_stderr=True, overwrite_output=True)
                )
                
                # Применяем фильтры к каждому каналу
                for channel in [left_channel, right_channel]:
                    processed_channel = channel.replace(".wav", "_processed.wav")
                    (
                        ffmpeg.input(channel)
                        .filter("highpass", f=120)  # Используем 'f' вместо 'freq'
                        .filter("lowpass", f=2800)  # Используем 'f' вместо 'freq'
                        .filter("loudnorm", i=-23, tp=-3, lra=9)
                        .output(processed_channel)
                        .run(capture_stdout=True, capture_stderr=True, overwrite_output=True)
                    )
            except ffmpeg.Error as e:
                logging.error("Ошибка в ffmpeg:")
                logging.error(e.stderr.decode())
                raise
            
            return {
                "left_channel": left_channel.replace(".wav", "_processed.wav"),
                "right_channel": right_channel.replace(".wav", "_processed.wav")
            }

        # Функция для транскрибации с помощью WhisperX
        def transcribe_audio(input_file, model_name="large-v2"):
            device = "cuda" if torch.cuda.is_available() else "cpu"
            model = whisperx.load_model(model_name, device=device)
            result = model.transcribe(input_file)
            return result["segments"]

        # Основной процесс
        try:
            filename = os.path.splitext(os.path.basename(audio_file))[0]
            metadata = parse_filename(filename)

            # Получаем длительность аудио
            duration = get_audio_duration(audio_file)
            if duration is not None:
                minutes = int(duration // 60)
                seconds = int(duration % 60)
                metadata["duration"] = f"{minutes} мин {seconds} сек"
            
            processed_files = preprocess_audio(audio_file, TEMP_DIR)
            
            # Транскрибация
            left_transcription = transcribe_audio(processed_files["left_channel"])
            right_transcription = transcribe_audio(processed_files["right_channel"])
            
            # Форматируем вывод
            left_formatted = format_transcription(left_transcription, "M")
            right_formatted = format_transcription(right_transcription, "K")
            full_transcription = left_formatted + "\n" + right_formatted
            
            # Сохраняем результаты
            metadata_file = os.path.join(TEMP_DIR, f"{filename}_metadata.json")
            with open(metadata_file, "w", encoding=ENCODING_TEXT) as f:
                json.dump(metadata, f, indent=4, default=str)
            
            transcription_file = os.path.join(TEMP_DIR, f"{filename}_transcription.txt")
            with open(transcription_file, "w", encoding=ENCODING_TEXT) as f:
                f.write(full_transcription)
            
            logging.info(f"Метаданные сохранены в '{metadata_file}'.")
            logging.info(f"Транскрипция сохранена в '{transcription_file}'.")
        except Exception as e:
            logging.error(f"Ошибка: {e}")

        return metadata_file, transcription_file, processed_files
    
    def clean_temp_files(self):
        # Удаление временных файлов
        for root, dirs, files in os.walk(TEMP_DIR):
            for file in files:
                os.remove(os.path.join(root, file))
            for dirname in dirs:
                os.rmdir(os.path.join(root, dirname))

    def analyze_wav(self, channels: dict, return_plt=False):
        
        # Функция для анализа тональности с обработкой ошибок
        def safe_analyze_emotion(segment, sr):
            try:
                return analyze_emotion(segment, sr)
            except Exception as e:
                logging.error(f"Ошибка при анализе сегмента: {e}")
                return None  # Возвращаем None в случае ошибки

        # Функция для анализа тональности
        def analyze_emotion(audio_segment, sr=16000):
            inputs = feature_extractor(audio_segment, sampling_rate=sr, return_tensors="pt", padding=True)
            with torch.no_grad():
                logits = model(**inputs).logits
            predicted_class = torch.argmax(logits, dim=-1).item()
            return predicted_class

        # Функция для разделения аудио на сегменты
        def split_audio(audio, segment_length, sr):
            return [audio[i:i + segment_length] for i in range(0, len(audio), segment_length)]

        # Функция для дополнения сегмента нулями
        def pad_segment(segment, target_length):
            if len(segment) < target_length:
                padding = target_length - len(segment)
                segment = np.pad(segment, (0, padding), mode='constant')
            return segment

        # Функция для агрегации эмоций
        def aggregate_emotions(emotions):
            avg_emotion = np.mean(emotions)
            positive_count = sum(1 for e in emotions if e == 1)  # Пример: 1 — позитивная эмоция
            neutral_count = sum(1 for e in emotions if e == 0)   # Пример: 0 — нейтральная
            negative_count = sum(1 for e in emotions if e == 2)  # Пример: 2 — негативная
            return avg_emotion, positive_count, neutral_count, negative_count
        
        # Интерпретация результатов
        def interpret_emotion(avg_emotion):
            if avg_emotion < 0.5:
                return "Нейтральная с уклоном в негативную"
            elif 0.5 <= avg_emotion < 1.5:
                return "Позитивная"
            else:
                return "Негативная"

        # Загрузка модели и процессора
        model_name = 'superb/wav2vec2-base-superb-er'
        model = Wav2Vec2ForSequenceClassification.from_pretrained(model_name)
        feature_extractor = Wav2Vec2FeatureExtractor.from_pretrained(model_name)

        # Проверка наличия файлов
        left_channel_path = channels['left_channel']
        right_channel_path = channels['right_channel']

        if not os.path.exists(left_channel_path):
            raise FileNotFoundError(f"Файл {left_channel_path} не найден.")
        if not os.path.exists(right_channel_path):
            raise FileNotFoundError(f"Файл {right_channel_path} не найден.")
        

        # Загрузка предобработанных аудиофайлов с частотой дискретизации 16000 Гц
        y1, sr = librosa.load(left_channel_path, sr=16000)
        y2, sr = librosa.load(right_channel_path, sr=16000)

        # Сегментация (например, по 5 секунд)
        segment_length = 5 * sr  # 5 секунд в сэмплах
        segments_channel_1 = split_audio(y1, segment_length, sr)
        segments_channel_2 = split_audio(y2, segment_length, sr)

        # Минимальная длина сегмента (например, 1 секунда)
        min_segment_length = 1 * sr

        # Анализ тональности для каждого сегмента с сохранением номеров негативных сегментов
        emotions_channel_1 = []
        emotions_channel_2 = []
        negative_segments = []

        for i, segment in enumerate(segments_channel_1):
            segment = pad_segment(segment, min_segment_length)  # Дополняем сегмент
            emotion = safe_analyze_emotion(segment, sr)
            if emotion is not None:  # Игнорируем сегменты с ошибками
                emotions_channel_1.append(emotion)
                if emotion == 2:  # Негативная эмоция
                    negative_segments.append(("Менеджер", i))

        for i, segment in enumerate(segments_channel_2):
            segment = pad_segment(segment, min_segment_length)  # Дополняем сегмент
            emotion = safe_analyze_emotion(segment, sr)
            if emotion is not None:  # Игнорируем сегменты с ошибками
                emotions_channel_2.append(emotion)
                if emotion == 2:  # Негативная эмоция
                    negative_segments.append(("Клиент", i))


            
        # Агрегация для каждого канала
        avg_emotion_1, positive_1, neutral_1, negative_1 = aggregate_emotions(emotions_channel_1)
        avg_emotion_2, positive_2, neutral_2, negative_2 = aggregate_emotions(emotions_channel_2)

        # Общая агрегация для диалога
        overall_avg_emotion = (avg_emotion_1 + avg_emotion_2) / 2

        interpretation_1 = interpret_emotion(avg_emotion_1)
        interpretation_2 = interpret_emotion(avg_emotion_2)
        interpretation_overall = interpret_emotion(overall_avg_emotion)

        report_filename = os.path.join(TEMP_DIR, 'tonality_report.txt')
        plot_filename = os.path.join(TEMP_DIR, 'tonality_plot.png')

        # Сохранение результатов в файл
        with open(report_filename, "w", encoding=ENCODING_TEXT) as f:
            f.write(f"Интерпретация тональности менеджера: {interpretation_1}\n")
            f.write(f"Интерпретация тональности клиента: {interpretation_2}\n")
            f.write(f"Интерпретация общей тональности диалога: {interpretation_overall}\n")
            if negative_segments:
                f.write("Негативные сегменты обнаружены в следующих местах:\n")
                for speaker, segment_num in negative_segments:
                    start_time = segment_num * 5  # Каждый сегмент длится 5 секунд
                    end_time = start_time + 5
                    f.write(f"- {speaker}, сегмент {segment_num} ({start_time}-{end_time} секунды)\n")
            else:
                f.write("Негативные сегменты не обнаружены.\n")

        # Построение графика тональности
        segments = range(len(segments_channel_1))
        plt.figure(figsize=(12, 6))
        plt.plot(segments, emotions_channel_1, label='Менеджер', marker='o', color='blue', linestyle='-', linewidth=2)
        plt.plot(segments, emotions_channel_2, label='Клиент', marker='x', color='green', linestyle='--', linewidth=2)

        # Подсветка негативных сегментов
        for speaker, segment_num in negative_segments:
            if speaker == "Менеджер":
                plt.scatter(segment_num, emotions_channel_1[segment_num], color='red', s=100, zorder=5, label='Негативный сегмент' if segment_num == 0 else "")
            else:
                plt.scatter(segment_num, emotions_channel_2[segment_num], color='red', s=100, zorder=5, label='Негативный сегмент' if segment_num == 0 else "")

        plt.xlabel('Номер сегмента', fontsize=12)
        plt.ylabel('Тональность', fontsize=12)
        plt.title('Тональность по сегментам (красные точки — негативные сегменты)', fontsize=14)
        plt.legend(fontsize=12)
        plt.grid(True, linestyle='--', alpha=0.7)
        plt.tight_layout()
        plt.savefig(plot_filename)  # Сохраняем график
        

        # Освобождение ресурсов
        del model, feature_extractor
        torch.cuda.empty_cache()

        # Вывод результатов
        logging.info(f"Анализ завершён. Результаты сохранены в файл '{report_filename}'.")
        logging.info(f"График тональности сохранён в файл '{plot_filename}'.")

        if return_plt:
            return report_filename, plot_filename, plt
        return report_filename, plot_filename
    
    def analyze_text(self, transcription_file):
        

        # Проверка существования файлов
        if not os.path.exists(FAISS_INDEX):
            raise FileNotFoundError(f"Файл {FAISS_INDEX} не найден. Проверьте путь и подключение Dataset.")
        if not os.path.exists(FAISS_METADATA):
            raise FileNotFoundError(f"Файл {FAISS_METADATA} не найден. Проверьте путь и подключение Dataset.")

        # Загрузка векторного индекса
        try:
            index = faiss.read_index(FAISS_INDEX)
            logging.info("Векторный индекс успешно загружен.")
        except Exception as e:
            logging.error(f"Ошибка при загрузке векторного индекса: {e}")

        # Загрузка метаданных
        try:
            with open(FAISS_METADATA, 'r') as f:
                metadata = json.load(f)
            logging.info("Метаданные успешно загружены.")
        except Exception as e:
            logging.error(f"Ошибка при загрузке метаданных: {e}")

        # Модель для создания эмбеддингов
        embedding_model = SentenceTransformer('all-MiniLM-L6-v2')

        # Функция для поиска релевантных чанков в базе знаний
        def search_relevant_chunks(query, top_k=3):
            query_embedding = embedding_model.encode([query])
            distances, indices = index.search(query_embedding, top_k)
            relevant_chunks = [metadata[i] for i in indices[0]]
            return relevant_chunks
        
        with open(transcription_file, 'r', encoding=ENCODING_TEXT) as f:
            transcript = f.read()

        # Поиск релевантных чанков для транскрипции
        relevant_chunks = search_relevant_chunks(transcript, top_k=10)  # Ищем 10 чанков

        # Объединение текста всех чанков
        criteria_text = "\n".join([chunk['text'] for chunk in relevant_chunks])  # Объединяем все чанки
        system_prompt = self.get_system_prompt(criteria_text)
        user_message = self.get_user_message(transcript)
        model = lms.llm('dab512/saiga_gemma2_9b-Q8_0-GGUF')
        respond = model.respond({'messages':[
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': user_message},
        ]})
        answer = respond.content
        out_filename = os.path.join(TEMP_DIR, "model_output.txt")
        with open(out_filename, "w", encoding=ENCODING_TEXT) as f:
            f.write(answer)

        return out_filename
    
    def create_report(self, metadata__path, transcription_path, tonality_report_path, model_output_path, tonality_plot_path):

        # Загрузка данных
        def load_data():
            # Метаданные
            with open(metadata__path, "r", encoding=ENCODING_TEXT) as f:
                metadata = json.load(f)
            
            # Текст транскрипции
            with open(transcription_path, "r", encoding=ENCODING_TEXT) as f:
                transcription = f.read()
            
            # Анализ эмоций
            with open(tonality_report_path, "r", encoding=ENCODING_TEXT) as f:
                tonality_report = f.read()
            
            # Ответ модели
            with open(model_output_path, "r", encoding=ENCODING_TEXT) as f:
                model_output = f.read()
            
            return metadata, transcription, tonality_report, model_output

        # Парсинг ответа модели
        def parse_model_output(model_output):
            lines = model_output.split("\n")
            criteria_scores = {}
            for line in lines:
                if ":" in line:
                    key, value = line.split(":", 1)
                    key = key.strip()
                    value = value.strip()
                    match = re.findall(r'\d', value[:5])
                    if match:
                        value = match[0]
                    if value.isdigit():
                        criteria_scores[key] = int(value)
            return criteria_scores

        # Преобразование оценок в проценты
        def convert_scores_to_percent(scores):
            conversion = {5: 100, 4: 75, 3: 50, 2: 25, 1: 0}
            return {k: conversion.get(v, 0) for k, v in scores.items()}

        # Учет критичных нарушений
        def apply_critical_violations(scores):
            if scores.get("Критичные нарушения", 0) > 0:
                for key in scores:
                    if key != "Критичные нарушения":
                        scores[key] = 0
            return scores
        def create_doc():
            metadata, transcription, tonality_report, model_output = load_data()
            doc = Document()
        
            # Заголовок
            doc.add_heading("Отчет по оценке работы менеджера", level=1)
            
            # Общая информация
            doc.add_paragraph(f"Дата: {metadata['date']}")
            doc.add_paragraph(f"Время: {metadata['time']}")
            doc.add_paragraph(f"Ф.И. менеджера: {metadata['employee_name']}")
            doc.add_paragraph(f"Номер телефона: {metadata['phone_number']}")
            doc.add_paragraph(f"Продолжительность звонка: {metadata['duration']}")
            
            # Текст транскрипции
            doc.add_heading("Транскрипция диалога", level=2)
            doc.add_paragraph(transcription)
            
            # Анализ модели
            doc.add_heading("Оценка работы менеджера", level=2)
            criteria_scores = parse_model_output(model_output)
            criteria_scores = convert_scores_to_percent(criteria_scores)
            criteria_scores = apply_critical_violations(criteria_scores)
            
            # Таблица с оценками (без строки "Критичные нарушения")
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Критерий"
            hdr_cells[1].text = "Оценка (%)"
            
            for criterion, score in criteria_scores.items():
                if criterion != "Критичные нарушения":  # Исключаем строку с критичными нарушениями
                    row_cells = table.add_row().cells
                    row_cells[0].text = criterion
                    row_cells[1].text = str(score)
            
            # Комментарий по критичным нарушениям
            critical_comment = None
            for line in model_output.split("\n"):
                if "Критичные нарушения" in line:
                    critical_comment = line.strip()
                    break
            
            if critical_comment:
                doc.add_paragraph("Комментарий по критическим нарушениям:")
                doc.add_paragraph(critical_comment)
            
            # Весь текст ответа модели
            doc.add_heading("Полный текст ответа модели", level=2)
            doc.add_paragraph(model_output)
            
            # Анализ эмоций
            doc.add_heading("Анализ эмоциональной оценки диалога", level=2)
            doc.add_paragraph(tonality_report)
            
            # Вставка графика эмоций
            doc.add_paragraph("График эмоциональной оценки:")
            doc.add_picture(tonality_plot_path, width=Inches(6))  # Ширина графика 6 дюймов
            
            # Сохранение документа
            report_path = os.path.join(TEMP_DIR, "manager_report.docx")
            doc.save(report_path)
            return report_path
        
        def create_zip(report_path):
            zip_path = os.path.join(OUTPUT_DIR, "report.zip")
            with zipfile.ZipFile(zip_path, 'w') as zipf:
                zipf.write(report_path, os.path.basename(report_path))
                zipf.write(tonality_plot_path, os.path.basename(tonality_plot_path))
            return zip_path
        
        report_path = create_doc()
        zip_path = create_zip(report_path)
        return zip_path